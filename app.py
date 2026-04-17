import streamlit as st
import pandas as pd
import plotly.express as px
import io
import unicodedata
import zipfile

try:
    from fpdf import FPDF
except ImportError:
    st.error("⚠️ A biblioteca 'fpdf2' não está instalada no ambiente do Streamlit.")
    st.info("Abra o terminal e digite: pip install fpdf2")
    st.stop()

try:
    from supabase import create_client, Client
except ImportError:
    st.error("⚠️ A biblioteca 'supabase' não está instalada no ambiente do Streamlit.")
    st.info("Abra o terminal e digite: pip install supabase")
    st.stop()

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    st.error("⚠️ A biblioteca 'python-docx' não está instalada.")
    st.info("Abra o terminal e digite: pip install python-docx")
    st.stop()

# ==========================================
# 1. CONFIGURAÇÃO INICIAL DA PÁGINA E ESTILOS
# ==========================================
st.set_page_config(page_title="Portal de Precificação", page_icon="🔒", layout="wide", initial_sidebar_state="expanded", menu_items={'Get Help': None, 'Report a bug': None, 'About': None})

COR_CABECALHO = "#7030A0"
COR_FUNDO_CLARO = "#E6E0EC"
COR_TEXTO_BRANCO = "#FFFFFF"
PALETA_GRAFICOS = ['#7030A0', '#9b59b6', '#3498db', '#1abc9c', '#f39c12', '#e74c3c']

st.markdown(f"""
    <style>
    .stApp {{ background-color: #F4F4F9; }}
    
    /* === OCULTAR ELEMENTOS PADRÃO DO STREAMLIT === */
    header {{ visibility: hidden !important; }}
    footer {{ visibility: hidden !important; }}
    .stAppDeployButton {{ display: none !important; }}
    #MainMenu {{ visibility: hidden !important; }}
    
    /* === Estilos da Tela de Login === */
    div[data-testid="stFormSubmitButton"] > button {{ background-color: {COR_CABECALHO} !important; color: white !important; border-radius: 8px !important; font-weight: bold !important; border: none !important; padding: 10px !important; }}
    div[data-testid="stFormSubmitButton"] > button:hover {{ background-color: #5a2680 !important; box-shadow: 0 4px 8px rgba(0,0,0,0.2) !important; }}
    
    /* === Estilos dos KPIs (Cartões) === */
    div[data-testid="stMetric"] {{
        background-color: #FFFFFF;
        border-left: 5px solid #7030A0;
        padding: 15px 20px;
        border-radius: 8px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        border: 1px solid #f0f0f0;
    }}
    
    /* === Correção dos Inputs === */
    div[data-baseweb="input"] > div, 
    div[data-baseweb="select"] > div {{
        background-color: #ffffff !important; 
        border: 1px solid #8e8e8e !important; 
        border-radius: 5px !important;
        transition: all 0.2s ease-in-out;
    }}
    div[data-baseweb="input"] > div:hover, 
    div[data-baseweb="select"] > div:hover {{ border-color: {COR_CABECALHO} !important; }}
    div[data-baseweb="input"] > div:focus-within, 
    div[data-baseweb="select"] > div:focus-within {{
        border: 2px solid {COR_CABECALHO} !important;
        box-shadow: 0 0 5px rgba(112, 48, 160, 0.2) !important;
    }}
    </style>
""", unsafe_allow_html=True)

# ==========================================
# 2. SISTEMA DE LOGIN
# ==========================================
try:
    SENHAS = st.secrets["senhas"]
except FileNotFoundError:
    st.error("⚠️ Arquivo secrets.toml não encontrado.")
    st.stop()
except KeyError:
    st.error("⚠️ A seção [senhas] não foi encontrada no secrets.toml.")
    st.stop()

if "usuario_logado" not in st.session_state:
    st.session_state["usuario_logado"] = None

if st.session_state["usuario_logado"] is None:
    st.write("<br><br><br>", unsafe_allow_html=True)
    col_esq, col_centro, col_dir = st.columns([1.5, 1.2, 1.5])
    
    with col_centro:
        st.markdown(f"""
        <div style="background-color: {COR_CABECALHO}; padding: 25px; border-radius: 10px 10px 0 0; text-align: center;">
            <h2 style="margin: 0; color: white;">SISTEMA DE PRECIFICAÇÃO</h2>
            <p style="color: #E6E0EC; margin-top: 5px; margin-bottom: 0;">Faça login para acessar seu painel</p>
        </div>
        """, unsafe_allow_html=True)
        
        with st.form("form_login"):
            st.write("")
            usuario_input = st.text_input("👤 Usuário")
            senha_input = st.text_input("🔑 Senha", type="password")
            st.write("")
            btn_entrar = st.form_submit_button("ENTRAR", use_container_width=True)
            
            if btn_entrar:
                if usuario_input in SENHAS and str(SENHAS[usuario_input]) == str(senha_input):
                    st.session_state["usuario_logado"] = usuario_input
                    st.rerun() 
                else:
                    st.error("❌ Usuário ou senha incorretos.")
    st.stop()

ID_CLIENTE = st.session_state["usuario_logado"]

# ==========================================
# 3. CONEXÃO COM A NUVEM E FUNÇÕES BASE
# ==========================================
@st.cache_resource
def iniciar_conexao():
    try:
        url = st.secrets["supabase"]["URL"]
        key = st.secrets["supabase"]["KEY"]
        return create_client(url, key)
    except Exception:
        return None

supabase: Client = iniciar_conexao()

def salvar_estado_nuvem():
    if supabase:
        dados = {
            "db_servicos": st.session_state.get("db_servicos", {}),
            "df_lista_equipamentos": st.session_state.get("df_lista_equipamentos", pd.DataFrame()).to_dict(orient="records"),
            "df_lista_insumos": st.session_state.get("df_lista_insumos", pd.DataFrame()).to_dict(orient="records"),
            "df_lista_taxas": st.session_state.get("df_lista_taxas", pd.DataFrame()).to_dict(orient="records"),
            "df_custos_categorias": {k: v.to_dict(orient="records") for k, v in st.session_state.get("df_custos_categorias", {}).items()},
            "protocolos": st.session_state.get("protocolos_db", [])
        }
        try:
            supabase.table("app_state").upsert({"cliente_id": ID_CLIENTE, "state_data": dados}).execute()
        except Exception as e:
            st.error(f"Erro ao salvar na nuvem: {e}")

def carregar_estado_nuvem():
    if supabase:
        try:
            res = supabase.table("app_state").select("state_data").eq("cliente_id", ID_CLIENTE).execute()
            if res.data:
                dados = res.data[0]["state_data"]
                st.session_state["db_servicos"] = dados.get("db_servicos", {})
                
                equip_data = dados.get("df_lista_equipamentos", [])
                insumos_data = dados.get("df_lista_insumos", [])
                taxas_data = dados.get("df_lista_taxas", [])

                st.session_state["df_lista_equipamentos"] = pd.DataFrame(equip_data) if equip_data else pd.DataFrame(columns=[
                    "Nome do equipamento", "Valor de aquisição (R$)", "Tempo de vida útil (anos)", 
                    "Capacidade de Aplicações / dia (R$)", "Aplicações (média diária)", "Custo anual de manutenção (R$)"
                ])

                st.session_state["df_lista_insumos"] = pd.DataFrame(insumos_data) if insumos_data else pd.DataFrame(columns=[
                    "Material", "qt", "valor"
                ])

                st.session_state["df_lista_taxas"] = pd.DataFrame(taxas_data) if taxas_data else pd.DataFrame(columns=[
                    "Taxa", "Porcentagem (%)"
                ])

                custos_raw = dados.get("df_custos_categorias", {})
                st.session_state["df_custos_categorias"] = {k: pd.DataFrame(v) for k, v in custos_raw.items()}
                
                st.session_state["protocolos_db"] = dados.get("protocolos", [])
                return True
        except Exception as e:
            st.error(f"Erro ao carregar dados da nuvem: {e}")
    return False

def df_maquinas_padrao(): return pd.DataFrame(columns=["nome", "custo"])
def df_insumos_padrao(): return pd.DataFrame(columns=["Material", "QT", "Preço (R$)"])
def df_outros_custos_padrao(): return pd.DataFrame(columns=["Tipo", "Descrição", "Valor (R$)", "Custo Fixo/Operação"])

def inicializar_padroes_caso_vazio():
    st.session_state["db_servicos"] = {}
    st.session_state["df_lista_equipamentos"] = pd.DataFrame(columns=["Nome do equipamento", "Valor de aquisição (R$)", "Tempo de vida útil (anos)", "Capacidade de Aplicações / dia (R$)", "Aplicações (média diária)", "Custo anual de manutenção (R$)"])
    st.session_state["df_lista_insumos"] = pd.DataFrame(columns=["Material", "qt", "valor"])
    st.session_state["df_lista_taxas"] = pd.DataFrame(columns=["Taxa", "Porcentagem (%)"])
    
    st.session_state["df_custos_categorias"] = {
        "1. Despesa com pessoal": pd.DataFrame([
            {"ÍTEM": "Total da folha de pagamento clt (com 13º)", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Despesas com alimentação e transporte", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Gratificações", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Uniformes", "MENSAL (R$)": 0.0}
        ]),
        "4. Despesas estrutura e de consumo": pd.DataFrame([
            {"ÍTEM": "Aluguel", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Energia Elétrica", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Água", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Internet", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Material de limpeza", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Material para copa / Experiencia do cliente", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Material de escritório (papel maca)", "MENSAL (R$)": 0.0},
            {"ÍTEM": "IPTU", "MENSAL (R$)": 0.0}
        ]),
        "8. Marketing e vendas": pd.DataFrame([
            {"ÍTEM": "Agência", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Eventos", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Tráfego Pago", "MENSAL (R$)": 0.0}
        ]),
        "Despesas Administrativas": pd.DataFrame([
            {"ÍTEM": "Contador", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Outras Consultorias", "MENSAL (R$)": 0.0}
        ]),
        "Despesas com TI": pd.DataFrame([
            {"ÍTEM": "Sistemas de gestão", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Chat bot / recursos de automação", "MENSAL (R$)": 0.0}
        ]),
        "Manutenção e Conservação": pd.DataFrame([
            {"ÍTEM": "Elevadores", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Coleta de lixo hospitalar", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Manutenção de equipamentos", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Pinturas", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Jardins", "MENSAL (R$)": 0.0},
            {"ÍTEM": "Extintores", "MENSAL (R$)": 0.0}
        ]),
        "Despesas Bancárias": pd.DataFrame([
            {"ÍTEM": "Taxa administrativa de contas", "MENSAL (R$)": 0.0}
        ]),
        "Seguros": pd.DataFrame([
            {"ÍTEM": "Seguros do estabelecimento", "MENSAL (R$)": 0.0}
        ])
    }
    st.session_state["protocolos_db"] = []
    salvar_estado_nuvem()

if "dados_carregados" not in st.session_state or st.session_state["dados_carregados"] == False:
    carregou_nuvem = carregar_estado_nuvem()
    if not carregou_nuvem:
        inicializar_padroes_caso_vazio()
    st.session_state["dados_carregados"] = True

if "dias_uteis_eq" not in st.session_state:
    st.session_state["dias_uteis_eq"] = 22.0

def carregar_servico_para_estado(nome_servico):
    if nome_servico not in st.session_state["db_servicos"]: return
    dados = st.session_state["db_servicos"][nome_servico]

    st.session_state["servico_atual"] = nome_servico
    st.session_state["tempo_min"] = float(dados.get("tempo_min", 60.0))
    st.session_state["repasse_fixo"] = float(dados.get("repasse_fixo", 0.0))
    st.session_state["repasse_percentual"] = float(dados.get("repasse_percentual", 0.0))
    st.session_state["custo_aluguel"] = float(dados.get("custo_aluguel", 0.0))
    
    st.session_state["df_ficha_maquinas"] = pd.DataFrame(dados.get("maquinas", [])) if dados.get("maquinas") else df_maquinas_padrao()
    st.session_state["df_ficha_insumos"] = pd.DataFrame(dados.get("insumos", [])) if dados.get("insumos") else df_insumos_padrao()
    st.session_state["df_ficha_outros_custos"] = pd.DataFrame(dados.get("outros_custos", [])) if dados.get("outros_custos") else df_outros_custos_padrao()
    
    taxas = dados.get("taxas", {})
    st.session_state["taxa_comissao"] = float(taxas.get("comissao", 0.0))
    st.session_state["cenario_cartao"] = taxas.get("cenario_cartao", "Crédito 1x")
    st.session_state["tipo_imposto"] = taxas.get("tipo_imposto", "Simples Nacional")
    st.session_state["aliquota_imposto"] = float(taxas.get("aliquota_imposto", 6.0))
    
    st.session_state["preco_escolhido"] = float(dados.get("preco_escolhido", 0.0))
    st.session_state.setdefault("indireto", "Sim")
    st.session_state.setdefault("valor_hora", 48.14)

def inicializar_estado_ficha():
    lista_nomes_servicos = list(st.session_state.get("db_servicos", {}).keys())
    if not lista_nomes_servicos:
        st.session_state.setdefault("servico_atual", "")
        st.session_state.setdefault("tempo_min", 60.0)
        st.session_state.setdefault("repasse_fixo", 0.0)
        st.session_state.setdefault("repasse_percentual", 0.0)
        st.session_state.setdefault("custo_aluguel", 0.0)
        st.session_state.setdefault("df_ficha_maquinas", df_maquinas_padrao())
        st.session_state.setdefault("df_ficha_insumos", df_insumos_padrao())
        st.session_state.setdefault("df_ficha_outros_custos", df_outros_custos_padrao())
        st.session_state.setdefault("taxa_comissao", 0.0)   
        st.session_state.setdefault("cenario_cartao", "Crédito 1x")
        st.session_state.setdefault("tipo_imposto", "Simples Nacional")
        st.session_state.setdefault("aliquota_imposto", 6.0)     
        st.session_state.setdefault("preco_escolhido", 0.0)
        st.session_state.setdefault("indireto", "Sim")
        st.session_state.setdefault("valor_hora", 48.14)
        return

    primeiro_servico = lista_nomes_servicos[0]
    carregar_servico_para_estado(primeiro_servico)

if "servico_atual" not in st.session_state:
    inicializar_estado_ficha()

# ==========================================
# GERAÇÃO DE PDF
# ==========================================
def gerar_pdf_ficha_tecnica(nome_servico, preco, custo_total, lucro, margem, impostos, taxa_cartao, comissao, repasse_med, df_maq, df_ins, df_outros):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    
    pdf.set_fill_color(112, 48, 160)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 10, f" FICHA TECNICA: {nome_servico.upper()}", 0, 1, 'C', fill=True)
    pdf.ln(5)
    
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 8, "1. RESUMO FINANCEIRO", 0, 1)
    pdf.set_font("Arial", '', 11)
    pdf.cell(50, 8, f"Preco de Venda:", 0, 0)
    pdf.cell(0, 8, f"R$ {preco:,.2f}", 0, 1)
    pdf.cell(50, 8, f"Custo Total:", 0, 0)
    pdf.cell(0, 8, f"R$ {custo_total:,.2f}", 0, 1)
    pdf.cell(50, 8, f"Deducoes (Imp/Tx/Com):", 0, 0)
    pdf.cell(0, 8, f"R$ {(impostos + taxa_cartao + comissao):,.2f}", 0, 1)
    pdf.cell(50, 8, f"Repasse Medico:", 0, 0)
    pdf.cell(0, 8, f"R$ {repasse_med:,.2f}", 0, 1)
    pdf.cell(50, 8, f"Lucro Liquido:", 0, 0)
    pdf.cell(0, 8, f"R$ {lucro:,.2f}  |  Margem: {margem*100:.1f}%", 0, 1)
    pdf.ln(5)

    def desenhar_tabela(titulo, df, colunas_mostrar):
        if not df.empty:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, titulo, 0, 1)
            pdf.set_font("Arial", 'B', 10)
            
            largura_col = 190 / len(colunas_mostrar)
            for col in colunas_mostrar:
                pdf.cell(largura_col, 8, str(col), border=1)
            pdf.ln()
            
            pdf.set_font("Arial", '', 10)
            for _, row in df.iterrows():
                for col in colunas_mostrar:
                    texto = str(row[col])[:30] 
                    pdf.cell(largura_col, 8, texto, border=1)
                pdf.ln()
            pdf.ln(5)

    desenhar_tabela("2. MAQUINAS E EQUIPAMENTOS", df_maq, ["nome", "custo"])
    desenhar_tabela("3. MATERIAIS E INSUMOS", df_ins, ["Material", "QT", "Preço (R$)"])
    if not df_outros.empty:
        desenhar_tabela("4. OUTROS CUSTOS", df_outros, ["Tipo", "Valor (R$)"])

    return bytes(pdf.output())

# ==========================================
# MENU LATERAL & ROTEAMENTO DE PÁGINAS
# ==========================================
with st.sidebar:    
    st.markdown(f"👤 Logado como: **{ID_CLIENTE}**")
    if st.button("🚪 Sair (Logout)", use_container_width=True):
        st.session_state.clear() 
        st.rerun()
    
    st.divider()
    
    modulo_selecionado = st.radio(
        "Navegue pelas ferramentas:",
        [
            "0. Início (Onboarding)",
            "1. Ficha Técnica (Precificação)",
            "2. Custos Fixos e Hora Clínica",
            "3. Registro de Equipamentos",
            "4. Insumos e Materiais",
            "5. Impostos e Taxas",
            "6. Protocolos (Jornadas) 🚀"
        ]
    )
    st.divider()
    if supabase:
        st.success(f"☁️ Nuvem Ativa")

# ==========================================
# MÓDULOS DE RENDERIZAÇÃO
# ==========================================
def cabecalho_padrao(titulo):
    st.markdown(f"""
    <div style="background-color: {COR_CABECALHO}; padding: 15px; border-radius: 8px; text-align: center; margin-bottom: 20px;">
        <h2 style="margin: 0; color: {COR_TEXTO_BRANCO};">{titulo}</h2>
    </div>
    """, unsafe_allow_html=True)


def render_onboarding():
    cabecalho_padrao("🚀 BEM-VINDO AO SEU PORTAL DE PRECIFICAÇÃO")
    st.markdown("Este painel foi criado para centralizar sua inteligência de preços. Assista aos tutoriais abaixo para configurar sua plataforma corretamente.")
    st.subheader("📺 Tutoriais em Vídeo")
    with st.expander("▶️ Como configurar os Custos Fixos"):
        st.info("Link: https://youtu.be/DXnuXyFi3V8")
    with st.expander("▶️ Como preencher a Ficha Técnica"):
        st.info("Link: youtube.com/...")
    with st.expander("▶️ Criando Protocolos e Propostas"):
        st.info("Link: youtube.com/...")


def render_ficha_tecnica():
    with st.sidebar:
        st.header("⚙️ Parâmetros Globais")
        st.selectbox("Considerar custo de hora clínica como indireto?", ["Sim", "Não"], key="indireto")
        st.number_input("Valor da Hora Clínica (R$)", min_value=0.0, step=1.0, format="%.2f", key="valor_hora")

    cabecalho_padrao("PORTAL DE PRECIFICAÇÃO CLÍNICA")

    lista_nomes_servicos = list(st.session_state["db_servicos"].keys())
    if not lista_nomes_servicos:
        st.warning("Nenhum serviço cadastrado. Vá na aba 'Gerenciar Serviços' e crie um novo.")
        with st.form("form_primeiro_servico"):
            n_nome = st.text_input("Criar Primeiro Serviço:")
            if st.form_submit_button("Criar"):
                if n_nome:
                    st.session_state["db_servicos"][n_nome] = {
                        "tempo_min": 60.0, "maquinas": [], "repasse_fixo": 0.0, "repasse_percentual": 0.0, "custo_aluguel": 0.0, "insumos": [], "outros_custos": [],
                        "taxas": {"comissao": 0.0, "cenario_cartao": "Crédito 1x", "tipo_imposto": "Simples Nacional", "aliquota_imposto": 6.0},
                        "preco_escolhido": 0.0
                    }
                    carregar_servico_para_estado(n_nome)
                    salvar_estado_nuvem()
                    st.rerun()
        return

    if "tempo_min" not in st.session_state or st.session_state.get("servico_atual") not in lista_nomes_servicos:
        carregar_servico_para_estado(lista_nomes_servicos[0])

    col_sel1, col_sel2 = st.columns([3, 1])
    with col_sel1:
        st.markdown("**🔍 Selecione o Serviço para Edição:**")
        st.selectbox(
            "Filtro", options=lista_nomes_servicos,
            index=lista_nomes_servicos.index(st.session_state["servico_atual"]) if st.session_state["servico_atual"] in lista_nomes_servicos else 0,
            key="combo_servico", on_change=lambda: carregar_servico_para_estado(st.session_state["combo_servico"]),
            label_visibility="collapsed"
        )
    with col_sel2:
        st.write("")
        if st.button("💾 Salvar na Nuvem", type="primary", use_container_width=True):
            st.session_state["db_servicos"][st.session_state["servico_atual"]] = {
                "tempo_min": st.session_state["tempo_min"],
                "maquinas": st.session_state.get("df_ficha_maquinas", df_maquinas_padrao()).to_dict(orient="records"),
                "repasse_fixo": st.session_state["repasse_fixo"],
                "repasse_percentual": st.session_state.get("repasse_percentual", 0.0),
                "custo_aluguel": st.session_state["custo_aluguel"],
                "insumos": st.session_state.get("df_ficha_insumos", df_insumos_padrao()).to_dict(orient="records"),
                "outros_custos": st.session_state.get("df_ficha_outros_custos", df_outros_custos_padrao()).to_dict(orient="records"),
                "taxas": {
                    "comissao": st.session_state["taxa_comissao"],
                    "cenario_cartao": st.session_state["cenario_cartao"],
                    "tipo_imposto": st.session_state["tipo_imposto"],
                    "aliquota_imposto": st.session_state["aliquota_imposto"]
                },
                "preco_escolhido": st.session_state["preco_escolhido"]
            }
            salvar_estado_nuvem()
            st.success("Salvo com sucesso!")

    st.write("")
    
    # -------------------------------
    # MATEMÁTICA E VARIÁVEIS
    # -------------------------------
    tempo_min = st.session_state.get("tempo_min", 60.0)
    valor_hora = st.session_state.get("valor_hora", 48.14)
    custo_execucao = (tempo_min / 60) * valor_hora if st.session_state.get("indireto") == "Sim" else 0.0
    
    df_maq = st.session_state.get("df_ficha_maquinas", df_maquinas_padrao())
    custo_maquinas = pd.to_numeric(df_maq["custo"], errors="coerce").fillna(0.0).sum() if not df_maq.empty else 0.0
    
    custo_aluguel_hora = st.session_state.get("custo_aluguel", 0.0)
    custo_aluguel_calculado = custo_aluguel_hora * (tempo_min / 60)
    
    repasse_fixo = st.session_state.get("repasse_fixo", 0.0)
    repasse_percentual = st.session_state.get("repasse_percentual", 0.0)
    
    df_ins = st.session_state.get("df_ficha_insumos", df_insumos_padrao())
    custo_insumos = (pd.to_numeric(df_ins["QT"], errors="coerce").fillna(0.0) * pd.to_numeric(df_ins["Preço (R$)"], errors="coerce").fillna(0.0)).sum() if not df_ins.empty else 0.0

    df_outros = st.session_state.get("df_ficha_outros_custos", df_outros_custos_padrao())
    custo_outros = pd.to_numeric(df_outros["Valor (R$)"], errors="coerce").fillna(0.0).sum() if not df_outros.empty else 0.0
    
    custo_total_servico = custo_execucao + custo_maquinas + custo_aluguel_calculado + repasse_fixo + custo_insumos + custo_outros
    
    df_taxas_globais = st.session_state.get("df_lista_taxas", pd.DataFrame())
    cenario_atual = st.session_state.get("cenario_cartao", "Crédito 1x")
    taxa_cartao_pct = 0.0
    if not df_taxas_globais.empty and cenario_atual in df_taxas_globais["Taxa"].values:
        taxa_cartao_pct = float(df_taxas_globais[df_taxas_globais["Taxa"] == cenario_atual]["Porcentagem (%)"].iloc[0])

    taxa_comissao_pct = st.session_state.get("taxa_comissao", 0.0)
    taxa_imposto_pct = st.session_state.get("aliquota_imposto", 6.0)
    preco_escolhido = st.session_state.get("preco_escolhido", 0.0)
    
    valor_imposto = preco_escolhido * (taxa_imposto_pct / 100)
    valor_taxa_cartao = preco_escolhido * (taxa_cartao_pct / 100)
    valor_comissao = preco_escolhido * (taxa_comissao_pct / 100)
    
    resultado_liquido = preco_escolhido - valor_imposto - valor_taxa_cartao - valor_comissao
    valor_repasse_medico = resultado_liquido * (repasse_percentual / 100)
    
    lucro = resultado_liquido - valor_repasse_medico - custo_total_servico
    pct_lucro = (lucro / preco_escolhido) if preco_escolhido > 0 else 0.0
    total_impostos_taxas = valor_imposto + valor_taxa_cartao + valor_comissao

    taxas_percentual_total = (taxa_comissao_pct + taxa_imposto_pct + taxa_cartao_pct) / 100
    divisor_break_even = (1 - taxas_percentual_total) * (1 - (repasse_percentual / 100))
    preco_sugerido_break_even = custo_total_servico / divisor_break_even if divisor_break_even > 0 else 0.0

    def cb_criar_servico():
        nome = st.session_state.get("input_novo_nome")
        if nome and nome not in st.session_state["db_servicos"]:
            st.session_state["db_servicos"][nome] = {
                "tempo_min": 60.0, "maquinas": [], "repasse_fixo": 0.0, "repasse_percentual": 0.0, "custo_aluguel": 0.0, "insumos": [], "outros_custos": [],
                "taxas": {"comissao": 0.0, "cenario_cartao": "Crédito 1x", "tipo_imposto": "Simples Nacional", "aliquota_imposto": 6.0},
                "preco_escolhido": 0.0
            }
            carregar_servico_para_estado(nome)
            salvar_estado_nuvem()
            st.session_state["input_novo_nome"] = ""

    def cb_renomear_servico():
        antigo = st.session_state.get("sel_ren_serv")
        novo = st.session_state.get("input_ren_serv")
        if novo and novo != antigo:
            novo_dict = {}
            for k, v in st.session_state["db_servicos"].items():
                if k == antigo: novo_dict[novo] = v
                else: novo_dict[k] = v
            st.session_state["db_servicos"] = novo_dict
            if st.session_state.get("servico_atual") == antigo:
                st.session_state["servico_atual"] = novo
            salvar_estado_nuvem()
            st.session_state["input_ren_serv"] = ""

    def cb_excluir_servico():
        remover = st.session_state.get("sel_rem_serv")
        if remover in st.session_state["db_servicos"]:
            del st.session_state["db_servicos"][remover]
            if st.session_state.get("servico_atual") == remover:
                st.session_state["servico_atual"] = ""
                if st.session_state["db_servicos"]: 
                    carregar_servico_para_estado(list(st.session_state["db_servicos"].keys())[0])
            salvar_estado_nuvem()

    def cb_excluir_todos():
        st.session_state["db_servicos"] = {}
        st.session_state["servico_atual"] = ""
        salvar_estado_nuvem()

    tab_dash, tab_custos, tab_precificacao, tab_gerenciar = st.tabs([
        "📊 Dashboard e Resumo", 
        "⚙️ Estrutura e Custos", 
        "💲 Precificação e Taxas",
        "🛠️ Gerenciar Serviços"
    ])

    with tab_dash:
        c_title, c_btn = st.columns([3, 1])
        c_title.markdown(f"<h4 style='color: {COR_CABECALHO};'>Resumo da Ficha: {st.session_state.get('servico_atual', '')}</h4>", unsafe_allow_html=True)
        
        with c_btn:
            pdf_bytes_unico = gerar_pdf_ficha_tecnica(
                st.session_state.get('servico_atual', ''), preco_escolhido, custo_total_servico, lucro, pct_lucro, 
                valor_imposto, valor_taxa_cartao, valor_comissao, valor_repasse_medico, df_maq, df_ins, df_outros
            )
            st.download_button(label="📄 Baixar PDF da Ficha", data=pdf_bytes_unico, file_name=f"Ficha_{st.session_state.get('servico_atual', '')}.pdf", mime="application/pdf", use_container_width=True)

        st.write("") 
        
        kpi1, kpi2, kpi3, kpi4 = st.columns(4)
        kpi1.metric("1. Preço de Venda", f"R$ {preco_escolhido:,.2f}")
        kpi2.metric("2. Deduções (Impostos/Taxas)", f"R$ {total_impostos_taxas:,.2f}")
        kpi3.metric("3. Resultado Líquido", f"R$ {resultado_liquido:,.2f}", help="O que entra no caixa após as taxas de venda.")
        kpi4.metric("4. Repasse Médico", f"R$ {valor_repasse_medico:,.2f}")

        st.write("") 
        
        sub_kpi1, sub_kpi2, sub_kpi3, sub_kpi4 = st.columns(4)
        sub_kpi1.metric("5. Custo Operacional", f"R$ {custo_total_servico:,.2f}", help="Soma de hora clínica, máquinas, insumos, aluguéis e repasses fixos.")
        sub_kpi2.metric("6. Lucro Líquido", f"R$ {lucro:,.2f}")
        sub_kpi3.metric("7. Margem (%)", f"{pct_lucro*100:.1f}%")
        sub_kpi4.metric("8. Preço Sugerido", f"R$ {preco_sugerido_break_even:,.2f}", help="Ponto de Equilíbrio: Preço exato para ter Lucro Zero.")

        st.write("") 

        with st.expander("📝 Entenda como cada valor foi calculado (Memorial de Cálculo)", expanded=True):
            st.markdown("### Memória de Cálculo")
            st.markdown("Acompanhe o caminho do dinheiro, do faturamento bruto até o lucro final:")
            
            c_mem1, c_mem2 = st.columns(2)
            with c_mem1:
                st.write("**1. Deduções sobre o Preço Bruto:**")
                st.latex(rf"Impostos ({taxa_imposto_pct}\%) = {valor_imposto:.2f}")
                st.latex(rf"Taxa\ Cartão ({taxa_cartao_pct}\%) = {valor_taxa_cartao:.2f}")
                st.latex(rf"Comissão ({taxa_comissao_pct}\%) = {valor_comissao:.2f}")
                st.info(f"**Resultado Líquido:** R$ {resultado_liquido:.2f} (O que sobra após taxas de venda)")

            with c_mem2:
                st.write("**2. Divisão do Resultado Líquido:**")
                st.latex(rf"Repasse\ ao\ Médico ({repasse_percentual}\%) = {valor_repasse_medico:.2f}")
                st.latex(rf"Custo\ Operacional = {custo_total_servico:.2f}")
                st.success(f"**Lucro Final:** R$ {lucro:.2f} (O que sobra no caixa da clínica)")
            
            st.markdown("---")
            st.write("**Fórmulas de Eficiência Financeira:**")
            st.latex(r"Margem\ (\%) = \left( \frac{Lucro\ Final}{Preço\ de\ Venda} \right) \times 100")
            st.latex(r"Preço\ Sugerido = \frac{Custo\ Operacional}{(1 - Taxas\%) \times (1 - Repasse\%)}")
            st.caption("💡 O **Preço Sugerido** (Ponto de Equilíbrio) faz o caminho inverso para descobrir quanto você precisa cobrar para cobrir todos os custos e pagar o médico, ficando com lucro zero no final.")

        st.divider()

        # ==========================================
        # PAINEL COMPARATIVO AVANÇADO
        # ==========================================
        st.markdown(f"<h4 style='color: {COR_CABECALHO};'>📈 Comparativo de Portfólio</h4>", unsafe_allow_html=True)
        
        todas_opcoes_servicos = list(st.session_state["db_servicos"].keys())
        total_cadastrados = len(todas_opcoes_servicos)

        st.markdown("**🎛️ Painel de Filtros do Gráfico:**")
        col_f1, col_f2 = st.columns([4, 1])
        
        with col_f2:
            st.write("")
            selecionar_todos = st.toggle("Selecionar Todos", value=True)
            
        with col_f1:
            padrao_selecao = todas_opcoes_servicos if selecionar_todos else []
            servicos_selecionados = st.multiselect(
                "Selecione os serviços que deseja colocar no gráfico:", 
                options=todas_opcoes_servicos, 
                default=padrao_selecao,
                help="Clique no 'X' para remover um serviço, ou digite para buscar."
            )
            
        metricas_opcoes = [
            "Preço de Venda (R$)", "Preço Sugerido (R$)", "Custo Total (R$)", "Lucro Líquido (R$)", 
            "Taxas/Impostos Totais (R$)", "Impostos (R$)", "Taxa Cartão (R$)", "Comissão (R$)",
            "Repasse Médico (R$)", "Hora Clínica (R$)", "Máquinas e Equip. (R$)", "Materiais e Insumos (R$)", 
            "Outros Custos (R$)", "Aluguel (R$)", "Repasse Fixo (R$)"
        ]
        
        metricas_selecionadas = st.multiselect(
            "Selecione os indicadores que deseja visualizar nas barras:",
            options=metricas_opcoes,
            default=["Preço de Venda (R$)", "Custo Total (R$)", "Lucro Líquido (R$)"]
        )

        if servicos_selecionados and metricas_selecionadas:
            st.success(f"👁️ Exibindo dados de **{len(servicos_selecionados)}** de **{total_cadastrados}** serviços cadastrados.")
            
            dados_comp = []
            valor_hora_global = st.session_state.get("valor_hora", 48.14)
            usa_indireto = st.session_state.get("indireto") == "Sim"

            for serv_nome in servicos_selecionados:
                dados_s = st.session_state["db_servicos"][serv_nome]
                
                t_min_l = dados_s.get("tempo_min", 60.0)
                c_exec_l = (t_min_l / 60) * valor_hora_global if usa_indireto else 0.0
                df_m_l = pd.DataFrame(dados_s.get("maquinas", []))
                c_maq_l = pd.to_numeric(df_m_l["custo"], errors="coerce").fillna(0.0).sum() if not df_m_l.empty else 0.0
                df_i_l = pd.DataFrame(dados_s.get("insumos", []))
                c_ins_l = (pd.to_numeric(df_i_l["QT"], errors="coerce").fillna(0.0) * pd.to_numeric(df_i_l["Preço (R$)"], errors="coerce").fillna(0.0)).sum() if not df_i_l.empty else 0.0
                df_o_l = pd.DataFrame(dados_s.get("outros_custos", []))
                c_outros_s_l = pd.to_numeric(df_o_l["Valor (R$)"], errors="coerce").fillna(0.0).sum() if not df_o_l.empty else 0.0
                
                c_alu_hora_l = dados_s.get("custo_aluguel", 0.0)
                c_alu_l = c_alu_hora_l * (t_min_l / 60)
                c_rep_l = dados_s.get("repasse_fixo", 0.0)
                
                preco_l = dados_s.get("preco_escolhido", 0.0)
                p_rep_med_l = dados_s.get("repasse_percentual", 0.0)
                taxas_s_l = dados_s.get("taxas", {})
                t_com_l = taxas_s_l.get("comissao", 0.0)
                cenario_s_l = taxas_s_l.get("cenario_cartao", "Crédito 1x")
                t_car_l = 0.0
                if not df_taxas_globais.empty and cenario_s_l in df_taxas_globais["Taxa"].values:
                    t_car_l = float(df_taxas_globais[df_taxas_globais["Taxa"] == cenario_s_l]["Porcentagem (%)"].iloc[0])
                t_imp_l = taxas_s_l.get("aliquota_imposto", 6.0)
                
                v_imp_l = preco_l * (t_imp_l / 100)
                v_car_l = preco_l * (t_car_l / 100)
                v_com_l = preco_l * (t_com_l / 100)
                
                liq_temp = preco_l - v_com_l - v_car_l - v_imp_l
                v_rep_med_l = liq_temp * (p_rep_med_l / 100)
                
                c_tot_l = c_exec_l + c_maq_l + c_ins_l + c_alu_l + c_rep_l + c_outros_s_l
                lucro_s_l = liq_temp - v_rep_med_l - c_tot_l
                margem_s_l = (lucro_s_l / preco_l) if preco_l > 0 else 0.0
                taxas_totais_l = v_imp_l + v_car_l + v_com_l
                
                taxas_pct_totais_l = (t_com_l + t_imp_l + t_car_l) / 100
                div_break_l = (1 - taxas_pct_totais_l) * (1 - (p_rep_med_l / 100))
                preco_sug_l = c_tot_l / div_break_l if div_break_l > 0 else 0.0
                
                dados_comp.append({
                    "Serviço": serv_nome, 
                    "Preço de Venda (R$)": preco_l, 
                    "Preço Sugerido (R$)": preco_sug_l,
                    "Custo Total (R$)": c_tot_l, 
                    "Lucro Líquido (R$)": lucro_s_l,
                    "Taxas/Impostos Totais (R$)": taxas_totais_l,
                    "Impostos (R$)": v_imp_l,
                    "Taxa Cartão (R$)": v_car_l,
                    "Comissão (R$)": v_com_l,
                    "Repasse Médico (R$)": v_rep_med_l,
                    "Hora Clínica (R$)": c_exec_l,
                    "Máquinas e Equip. (R$)": c_maq_l,
                    "Materiais e Insumos (R$)": c_ins_l,
                    "Outros Custos (R$)": c_outros_s_l,
                    "Aluguel (R$)": c_alu_l,
                    "Repasse Fixo (R$)": c_rep_l,
                    "Margem": margem_s_l
                })
                
            df_comp = pd.DataFrame(dados_comp)
            df_comp = df_comp.sort_values(by="Lucro Líquido (R$)", ascending=True) 
            
            st.markdown("##### 💡 Insights da Seleção:")
            col_in1, col_in2, col_in3 = st.columns(3)
            servico_top = df_comp.iloc[-1] 
            lucro_medio = df_comp["Lucro Líquido (R$)"].mean()
            
            col_in1.metric("Serviço Mais Rentável", servico_top["Serviço"])
            col_in2.metric("Pico de Lucro", f"R$ {servico_top['Lucro Líquido (R$)']:,.2f}")
            col_in3.metric("Média de Lucro (Seleção)", f"R$ {lucro_medio:,.2f}")
            
            st.write("")

            df_melted = df_comp.melt(id_vars=["Serviço"], value_vars=metricas_selecionadas, var_name="Métrica", value_name="Valor")
            fig_comp = px.bar(
                df_melted, y="Serviço", x="Valor", color="Métrica", 
                barmode="group", orientation='h', color_discrete_sequence=px.colors.qualitative.Plotly
            )
            fig_comp.update_traces(texttemplate='R$ %{x:,.0f}', textposition='outside', textfont_size=11)
            max_val = df_melted["Valor"].max()
            altura_dinamica = max(400, len(df_comp) * 85) 
            
            fig_comp.update_layout(
                plot_bgcolor="rgba(0,0,0,0)", xaxis_title="Reais (R$)", yaxis_title="",
                margin=dict(t=10, b=50, l=150), legend_title="",
                legend=dict(orientation="h", yanchor="top", y=-0.15, xanchor="center", x=0.5),
                xaxis=dict(range=[0, max_val * 1.25]), height=altura_dinamica
            )
            st.plotly_chart(fig_comp, use_container_width=True)
            
            with st.expander("📄 Ver Tabela de Comparação Completa"):
                df_exibicao = df_comp.sort_values(by="Lucro Líquido (R$)", ascending=False)
                st.dataframe(df_exibicao.style.format({
                    "Preço de Venda (R$)": "R$ {:,.2f}", 
                    "Preço Sugerido (R$)": "R$ {:,.2f}",
                    "Custo Total (R$)": "R$ {:,.2f}", 
                    "Lucro Líquido (R$)": "R$ {:,.2f}",
                    "Taxas/Impostos Totais (R$)": "R$ {:,.2f}",
                    "Impostos (R$)": "R$ {:,.2f}",
                    "Taxa Cartão (R$)": "R$ {:,.2f}",
                    "Comissão (R$)": "R$ {:,.2f}",
                    "Repasse Médico (R$)": "R$ {:,.2f}",
                    "Hora Clínica (R$)": "R$ {:,.2f}",
                    "Máquinas e Equip. (R$)": "R$ {:,.2f}",
                    "Materiais e Insumos (R$)": "R$ {:,.2f}",
                    "Outros Custos (R$)": "R$ {:,.2f}",
                    "Aluguel (R$)": "R$ {:,.2f}",
                    "Repasse Fixo (R$)": "R$ {:,.2f}",
                    "Margem": "{:.1%}"
                }), use_container_width=True, hide_index=True)
        else:
            st.warning("⚠️ Selecione pelo menos um serviço e uma métrica para exibir o gráfico.")

        # ==========================================
        # EXPORTAÇÃO EM LOTE (ZIP COM TODOS OS PDFs)
        # ==========================================
        st.divider()
        st.markdown(f"<h4 style='color: {COR_CABECALHO};'>📦 Exportação em Lote</h4>", unsafe_allow_html=True)
        st.info("Selecione os serviços para gerar um arquivo ZIP contendo todas as Fichas Técnicas em PDF de uma só vez.")

        servicos_para_zip = st.multiselect(
            "Fichas para incluir no ZIP:",
            options=todas_opcoes_servicos,
            default=todas_opcoes_servicos,
            key="zip_multiselect"
        )

        if st.button("🗜️ Gerar Arquivo ZIP com as Fichas", type="primary", use_container_width=True):
            if servicos_para_zip:
                zip_buffer = io.BytesIO()
                
                with st.spinner("Gerando PDFs e compactando... Isso pode levar alguns segundos."):
                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                        for serv_nome in servicos_para_zip:
                            dados_s = st.session_state["db_servicos"][serv_nome]
                            
                            t_min_z = dados_s.get("tempo_min", 60.0)
                            v_hora_z = st.session_state.get("valor_hora", 48.14)
                            usa_ind_z = st.session_state.get("indireto") == "Sim"
                            c_exec_z = (t_min_z / 60) * v_hora_z if usa_ind_z else 0.0

                            df_m_z = pd.DataFrame(dados_s.get("maquinas", []))
                            c_maq_z = pd.to_numeric(df_m_z["custo"], errors="coerce").fillna(0.0).sum() if not df_m_z.empty else 0.0

                            df_i_z = pd.DataFrame(dados_s.get("insumos", []))
                            c_ins_z = (pd.to_numeric(df_i_z["QT"], errors="coerce").fillna(0.0) * pd.to_numeric(df_i_z["Preço (R$)"], errors="coerce").fillna(0.0)).sum() if not df_i_z.empty else 0.0

                            df_o_z = pd.DataFrame(dados_s.get("outros_custos", []))
                            c_outros_s_z = pd.to_numeric(df_o_z["Valor (R$)"], errors="coerce").fillna(0.0).sum() if not df_o_z.empty else 0.0

                            c_alu_hora_z = dados_s.get("custo_aluguel", 0.0)
                            c_alu_z = c_alu_hora_z * (t_min_z / 60)
                            c_rep_z = dados_s.get("repasse_fixo", 0.0)

                            preco_z = dados_s.get("preco_escolhido", 0.0)
                            p_rep_med_z = dados_s.get("repasse_percentual", 0.0)
                            taxas_s_z = dados_s.get("taxas", {})
                            t_com_z = taxas_s_z.get("comissao", 0.0)

                            cenario_s_z = taxas_s_z.get("cenario_cartao", "Crédito 1x")
                            t_car_z = 0.0
                            df_taxas_g = st.session_state.get("df_lista_taxas", pd.DataFrame())
                            if not df_taxas_g.empty and cenario_s_z in df_taxas_g["Taxa"].values:
                                t_car_z = float(df_taxas_g[df_taxas_g["Taxa"] == cenario_s_z]["Porcentagem (%)"].iloc[0])

                            t_imp_z = taxas_s_z.get("aliquota_imposto", 6.0)

                            c_tot_z = c_exec_z + c_maq_z + c_ins_z + c_alu_z + c_rep_z + c_outros_s_z
                            v_imp_z = preco_z * (t_imp_z / 100)
                            v_car_z = preco_z * (t_car_z / 100)
                            v_com_z = preco_z * (t_com_z / 100)
                            
                            liq_temp_z = preco_z - v_com_z - v_car_z - v_imp_z
                            v_rep_med_z = liq_temp_z * (p_rep_med_z / 100)
                            
                            lucro_s_z = liq_temp_z - v_rep_med_z - c_tot_z
                            margem_s_z = (lucro_s_z / preco_z) if preco_z > 0 else 0.0

                            pdf_bytes = gerar_pdf_ficha_tecnica(
                                serv_nome, preco_z, c_tot_z, lucro_s_z, margem_s_z, 
                                v_imp_z, v_car_z, v_com_z, v_rep_med_z, df_m_z, df_i_z, df_o_z
                            )

                            nome_arquivo_limpo = "".join(c for c in serv_nome if c.isalnum() or c in (' ', '-', '_')).rstrip()
                            zip_file.writestr(f"Ficha_{nome_arquivo_limpo}.pdf", pdf_bytes)

                zip_buffer.seek(0)
                st.success("✅ Arquivo ZIP gerado com sucesso!")
                st.download_button(
                    label="⬇️ CLIQUE AQUI PARA SALVAR O ZIP",
                    data=zip_buffer,
                    file_name="Todas_As_Fichas_Tecnicas.zip",
                    mime="application/zip",
                    use_container_width=True
                )
            else:
                st.warning("Selecione pelo menos um serviço para gerar o arquivo.")

    with tab_custos:
        st.subheader("Tempo e Repasses")
        c_t1, c_t2, c_t3 = st.columns(3)
        c_t1.number_input("TEMPO de execução (Minutos):", min_value=0.0, step=5.0, format="%.0f", key="tempo_min")
        c_t2.number_input("Aluguel Máquina (R$/Hora):", min_value=0.0, step=10.0, format="%.2f", key="custo_aluguel", help="O sistema calculará proporcional ao tempo da sessão.")
        c_t3.number_input("Repasse profissionais (Fixo R$):", min_value=0.0, step=10.0, format="%.2f", key="repasse_fixo")
        
        st.divider()
        st.subheader("👥 Outros Custos (Equipe, Feriados, Adicionais)")
        if not df_outros.empty:
            st.dataframe(df_outros.style.format({"Valor (R$)": "R$ {:.2f}"}), use_container_width=True, hide_index=True)
        else:
            st.info("Nenhum custo extra cadastrado (Ex: Enfermeiro, Instrumentador).")
        
        with st.expander("➕ Adicionar Outro Custo"):
            with st.form("form_outros_custos", clear_on_submit=True):
                col_oc1, col_oc2, col_oc3, col_oc4 = st.columns([2, 2, 2, 2])
                tipo_oc = col_oc1.selectbox("Tipo", ["Auxiliar Técnico", "Enfermeiro", "Instrumentador", "Segurança", "Taxa Feriado", "Outro (Livre)"])
                desc_oc = col_oc2.text_input("Descrição (se outro)")
                val_oc = col_oc3.number_input("Valor (R$)", min_value=0.0, step=10.0)
                freq_oc = col_oc4.selectbox("Cobrança", ["Por Operação", "Fixo Mensal (Rateio)"]) 
                
                if st.form_submit_button("Adicionar Custo"):
                    nome_final = desc_oc if tipo_oc == "Outro (Livre)" and desc_oc else tipo_oc
                    novo_oc = pd.DataFrame([{"Tipo": tipo_oc, "Descrição": nome_final, "Valor (R$)": float(val_oc), "Custo Fixo/Operação": freq_oc}])
                    st.session_state["df_ficha_outros_custos"] = pd.concat([df_outros, novo_oc], ignore_index=True)
                    st.rerun()

        st.divider()
        st.subheader("🖥️ Máquinas e Equipamentos Utilizados")
        if not df_maq.empty:
            st.dataframe(df_maq.style.format({"custo": "R$ {:.2f}"}), use_container_width=True, hide_index=True)
        else:
            st.info("Nenhuma máquina cadastrada para este serviço.")

        tab_add_m, tab_ren_m, tab_del_m = st.tabs(["➕ Adicionar", "✏️ Editar", "🗑️ Excluir"])
        with tab_add_m:
            df_eq_global = st.session_state.get("df_lista_equipamentos", pd.DataFrame())
            opcoes_eq = ["-- Digitar Manualmente --"] + df_eq_global["Nome do equipamento"].tolist() if not df_eq_global.empty else ["-- Digitar Manualmente --"]
            sel_eq = st.selectbox("Buscar Equipamento Cadastrado:", opcoes_eq, key="sel_add_maq_ficha")
            
            default_nome_eq = sel_eq if sel_eq != "-- Digitar Manualmente --" else ""
            default_custo_eq = 0.0
            if sel_eq != "-- Digitar Manualmente --":
                row = df_eq_global[df_eq_global["Nome do equipamento"] == sel_eq].iloc[0]
                montante = row["Valor de aquisição (R$)"] + (row["Tempo de vida útil (anos)"] * row.get("Custo anual de manutenção (R$)", 0.0))
                dep = montante / (row["Tempo de vida útil (anos)"] * 12) if row["Tempo de vida útil (anos)"] > 0 else 0
                dias_uteis = st.session_state.get("dias_uteis_eq", 22.0)
                default_custo_eq = dep / (row.get("Aplicações (média diária)", 1) * dias_uteis) if row.get("Aplicações (média diária)", 0) > 0 else 0

            with st.form("form_add_maq", clear_on_submit=True):
                c1, c2, c3 = st.columns([4, 2, 2])
                n_nome = c1.text_input("Nome da Máquina", value=default_nome_eq)
                n_custo = c2.number_input("Custo da Seção (R$)", value=float(default_custo_eq), min_value=0.0, step=10.0, format="%.2f")
                if c3.form_submit_button("Adicionar"):
                    if n_nome:
                        novo_reg = pd.DataFrame([{"nome": n_nome, "custo": float(n_custo)}])
                        st.session_state["df_ficha_maquinas"] = pd.concat([df_maq, novo_reg], ignore_index=True)
                        st.rerun()

        with tab_ren_m:
            if not df_maq.empty:
                c1, c2, c3 = st.columns([2, 2, 1])
                maq_renomear = c1.selectbox("Máquina atual:", options=df_maq["nome"].tolist(), key="sel_ren_maq_ficha")
                custo_atual_maq = df_maq[df_maq["nome"] == maq_renomear]["custo"].iloc[0]
                novo_nome_maq = c1.text_input("Mudar nome para:", value=maq_renomear, key="in_ren_maq_ficha")
                novo_custo_maq = c2.number_input("Mudar custo para (R$):", value=float(custo_atual_maq), min_value=0.0, step=10.0, format="%.2f", key="in_ren_custo_maq_ficha")
                st.write("")
                if c3.button("Atualizar", key="btn_salvar_maq", use_container_width=True):
                    if novo_nome_maq:
                        idx = df_maq.index[df_maq["nome"] == maq_renomear].tolist()[0]
                        df_maq.at[idx, "nome"] = novo_nome_maq
                        df_maq.at[idx, "custo"] = novo_custo_maq
                        st.session_state["df_ficha_maquinas"] = df_maq
                        st.rerun()

        with tab_del_m:
            if not df_maq.empty:
                c1, c2, c3 = st.columns([2, 1, 1])
                maq_remover = c1.selectbox("Remover máquina:", options=df_maq["nome"].tolist(), key="sel_rem_maq_ficha")
                st.write("")
                if c2.button("🗑️ Remover", key="btn_rem_maq", use_container_width=True):
                    st.session_state["df_ficha_maquinas"] = df_maq[df_maq["nome"] != maq_remover]
                    st.rerun()
                st.write("")
                if c3.button("⚠️ Excluir TODAS", key="btn_rem_todas_maq", type="primary", use_container_width=True):
                    st.session_state["df_ficha_maquinas"] = df_maquinas_padrao()
                    st.rerun()

        st.divider()
        st.subheader("💉 Materiais e Insumos Utilizados")
        if not df_ins.empty:
            st.dataframe(df_ins.style.format({"QT": "{:.2f}", "Preço (R$)": "R$ {:.3f}"}), use_container_width=True, hide_index=True)
        else:
            st.info("Nenhum insumo cadastrado para este serviço.")

        tab_add_i, tab_ren_i, tab_del_i = st.tabs(["➕ Adicionar", "✏️ Editar", "🗑️ Excluir"])
        with tab_add_i:
            df_ins_db = st.session_state.get("df_lista_insumos", pd.DataFrame())
            opcoes_ins = ["-- Digitar Manualmente --"] + df_ins_db["Material"].tolist() if not df_ins_db.empty else ["-- Digitar Manualmente --"]
            sel_ins = st.selectbox("Buscar Insumo Cadastrado:", opcoes_ins, key="sel_add_ins_ficha")
            
            default_mat = sel_ins if sel_ins != "-- Digitar Manualmente --" else ""
            default_preco = 0.0
            if sel_ins != "-- Digitar Manualmente --":
                row_ins = df_ins_db[df_ins_db["Material"] == sel_ins].iloc[0]
                qt_base = float(row_ins.get("qt", 1.0))
                qt_base = qt_base if qt_base > 0 else 1.0
                default_preco = float(row_ins.get("valor", 0.0)) / qt_base

            with st.form("form_add_ins_ficha", clear_on_submit=True):
                c1, c2, c3, c4 = st.columns([3, 2, 2, 2])
                n_mat = c1.text_input("Material", value=default_mat)
                n_qt = c2.number_input("Qtd a usar", value=1.00, min_value=0.01, step=1.0)
                n_preco = c3.number_input("Preço Un. (R$)", value=float(default_preco), min_value=0.0, step=0.10, format="%.3f")
                if c4.form_submit_button("Adicionar"):
                    if n_mat:
                        novo_reg = pd.DataFrame([{"Material": n_mat, "QT": float(n_qt), "Preço (R$)": float(n_preco)}])
                        st.session_state["df_ficha_insumos"] = pd.concat([df_ins, novo_reg], ignore_index=True)
                        st.rerun()

        with tab_ren_i:
            if not df_ins.empty:
                c1, c2, c3, c4 = st.columns([2, 1, 1, 1])
                ins_renomear = c1.selectbox("Insumo atual:", options=df_ins["Material"].tolist(), key="sel_ren_ins_ficha")
                row_atual = df_ins[df_ins["Material"] == ins_renomear].iloc[0]
                
                novo_nome_ins = c1.text_input("Mudar nome para:", value=ins_renomear, key="in_ren_ins_ficha")
                nova_qt_ins = c2.number_input("Mudar Qtd:", value=float(row_atual["QT"]), min_value=0.01, step=1.0, key="in_ren_qt_ins")
                novo_preco_ins = c3.number_input("Mudar Preço Un.:", value=float(row_atual["Preço (R$)"]), min_value=0.0, step=0.1, format="%.3f", key="in_ren_pr_ins")
                st.write("")
                if c4.button("Atualizar", key="btn_salvar_ins", use_container_width=True):
                    if novo_nome_ins:
                        idx = df_ins.index[df_ins["Material"] == ins_renomear].tolist()[0]
                        df_ins.at[idx, "Material"] = novo_nome_ins
                        df_ins.at[idx, "QT"] = nova_qt_ins
                        df_ins.at[idx, "Preço (R$)"] = novo_preco_ins
                        st.session_state["df_ficha_insumos"] = df_ins
                        st.rerun()

        with tab_del_i:
            if not df_ins.empty:
                c1, c2, c3 = st.columns([2, 1, 1])
                ins_remover = c1.selectbox("Remover insumo:", options=df_ins["Material"].tolist(), key="sel_rem_ins_ficha")
                st.write("")
                if c2.button("🗑️ Remover", key="btn_rem_ins", use_container_width=True):
                    st.session_state["df_ficha_insumos"] = df_ins[df_ins["Material"] != ins_remover]
                    st.rerun()
                st.write("")
                if c3.button("⚠️ Excluir TODOS", key="btn_rem_todos_ins", type="primary", use_container_width=True):
                    st.session_state["df_ficha_insumos"] = df_insumos_padrao()
                    st.rerun()

    with tab_precificacao:
        st.subheader("Configuração Fiscal e Repasses")
        col_f1, col_f2, col_f3 = st.columns(3)
        col_f1.selectbox("Modelo / Tipo de Imposto", ["Simples Nacional", "Lucro Presumido", "Lucro Real", "Isento/PF"], key="tipo_imposto")
        col_f2.number_input("Alíquota do Imposto (%)", min_value=0.0, step=0.1, format="%.2f", key="aliquota_imposto")
        col_f3.number_input("Repasse Médico (%)", min_value=0.0, step=1.0, format="%.2f", key="repasse_percentual", help="Calculado em cima do Valor Líquido (após impostos e taxas).")

        st.divider()
        opcoes_taxas = df_taxas_globais["Taxa"].tolist() if not df_taxas_globais.empty else ["Nenhuma taxa cadastrada"]
        
        col_t1, col_t2 = st.columns(2)
        col_t1.selectbox("Cenário de Pagamento Padrão", opcoes_taxas, key="cenario_cartao")
        col_t2.number_input("COMISSÃO Venda (%)", min_value=0.0, step=0.1, format="%.2f", key="taxa_comissao")

        st.markdown("### Preço de Venda")
        st.number_input("PREÇO DE TABELA FINAL (R$)", min_value=0.0, step=10.0, format="%.2f", key="preco_escolhido")
        st.caption(f"💡 **Dica do Sistema:** Para ter **LUCRO ZERO**, seu preço de tabela precisaria ser de pelo menos **R$ {preco_sugerido_break_even:,.2f}**.")

    with tab_gerenciar:
        st.info("Aqui você pode criar um serviço em branco, renomear um existente ou apagar serviços que não usa mais.")
        opcoes_servicos = list(st.session_state["db_servicos"].keys())
        
        tab_add_s, tab_ren_s, tab_del_s = st.tabs(["➕ Criar Novo Serviço", "✏️ Renomear Atual", "🗑️ Excluir"])

        with tab_add_s:
            c1, c2 = st.columns([2, 1])
            c1.text_input("Nome do Novo Serviço:", key="input_novo_nome")
            c2.write("")
            c2.button("Criar Serviço", on_click=cb_criar_servico, use_container_width=True)

        with tab_ren_s:
            c1, c2, c3 = st.columns([2, 2, 1])
            c1.selectbox("Serviço para renomear:", options=opcoes_servicos, key="sel_ren_serv")
            c2.text_input("Mudar para:", key="input_ren_serv")
            c3.write("")
            c3.button("Salvar Nome", on_click=cb_renomear_servico, use_container_width=True)

        with tab_del_s:
            c1, c2, c3 = st.columns([2, 1, 1])
            c1.selectbox("Serviço para excluir:", options=opcoes_servicos, key="sel_rem_serv")
            c2.write("")
            c2.button("🗑️ Remover", on_click=cb_excluir_servico, use_container_width=True)
            c3.write("")
            c3.button("⚠️ Excluir TODOS", on_click=cb_excluir_todos, type="primary", use_container_width=True)

def render_custos_fixos():
    cabecalho_padrao("GESTÃO DE CUSTOS FIXOS E HORA CLÍNICA")
    
    col_t1, col_t2 = st.columns([3, 1])
    with col_t2:
        if st.button("💾 Salvar na Nuvem", type="primary", use_container_width=True):
            salvar_estado_nuvem()
            st.success("Custos sincronizados!")

    def renderizar_categoria_dinamica(titulo, chave):
        with st.expander(titulo, expanded=False):
            df_atual = st.session_state["df_custos_categorias"][titulo]
            registros = df_atual.to_dict('records')

            if registros:
                for idx, reg in enumerate(registros):
                    c1, c2, c3 = st.columns([5, 3, 1])
                    novo_item = c1.text_input("Descrição", value=reg['ÍTEM'], key=f"edit_item_{chave}_{idx}", label_visibility="collapsed")
                    novo_valor = c2.number_input("Valor", value=float(reg['MENSAL (R$)']), min_value=0.0, step=10.0, format="%.2f", key=f"edit_val_{chave}_{idx}", label_visibility="collapsed")
                    
                    registros[idx]['ÍTEM'] = novo_item
                    registros[idx]['MENSAL (R$)'] = novo_valor
                    
                    if c3.button("🗑️", key=f"del_{chave}_{idx}"):
                        registros.pop(idx)
                        st.session_state["df_custos_categorias"][titulo] = pd.DataFrame(registros) if registros else pd.DataFrame(columns=["ÍTEM", "MENSAL (R$)"])
                        st.rerun()
                        
                st.session_state["df_custos_categorias"][titulo] = pd.DataFrame(registros) if registros else pd.DataFrame(columns=["ÍTEM", "MENSAL (R$)"])
            else:
                st.info("Nenhuma despesa cadastrada.")

            st.markdown("---")
            with st.form(f"form_add_{chave}", clear_on_submit=True):
                st.caption("Adicionar nova despesa:")
                c1, c2, c3 = st.columns([4, 2, 2])
                n_item = c1.text_input("Descrição da Despesa")
                n_valor = c2.number_input("Valor (R$)", min_value=0.0, step=10.0, format="%.2f")
                if c3.form_submit_button("➕ Adicionar"):
                    if n_item:
                        registros.append({"ÍTEM": n_item, "MENSAL (R$)": float(n_valor)})
                        st.session_state["df_custos_categorias"][titulo] = pd.DataFrame(registros)
                        st.rerun()

            total = sum(float(r["MENSAL (R$)"]) for r in registros) if registros else 0.0
            st.markdown(f"**Total da Categoria:** R$ {total:,.2f}")
            return total

    despesa_mensal_media = 0.0
    for i, categoria in enumerate(st.session_state["df_custos_categorias"].keys()):
        df_atual = st.session_state["df_custos_categorias"][categoria]
        tot_cat = pd.to_numeric(df_atual["MENSAL (R$)"], errors="coerce").fillna(0.0).sum() if not df_atual.empty else 0.0
        despesa_mensal_media += tot_cat

    horas_diarias = st.session_state.get("horas_diarias", 8.0)
    dias_semana = st.session_state.get("dias_semana", 5.0)
    qtd_salas = st.session_state.get("qtd_salas", 14.0)

    horas_semanais = horas_diarias * dias_semana
    horas_mensais = horas_semanais * 4.5
    despesa_anual = despesa_mensal_media * 12

    custo_hora_clinica = despesa_mensal_media / horas_mensais if horas_mensais > 0 else 0.0
    custo_dia_clinica = custo_hora_clinica * horas_diarias
    custo_hora_atendimento = custo_hora_clinica / qtd_salas if qtd_salas > 0 else 0.0

    tab_dash, tab_lanc, tab_cat = st.tabs(["📊 Dashboard Geral", "📝 Lançamentos e Hora Clínica", "🛠️ Criar Categorias"])

    with tab_dash:
        st.markdown(f"<h3 style='color: {COR_CABECALHO};'>Visão Geral dos Custos</h3>", unsafe_allow_html=True)
        kpi1, kpi2, kpi3, kpi4 = st.columns(4)
        kpi1.metric("Despesa Mensal", f"R$ {despesa_mensal_media:,.2f}")
        kpi2.metric("Despesa Anual", f"R$ {despesa_anual:,.2f}")
        kpi3.metric("Custo Hora Clínica", f"R$ {custo_hora_clinica:,.2f}")
        kpi4.metric("Custo Hora (por Sala)", f"R$ {custo_hora_atendimento:,.2f}")

        st.write("")
        st.write("")

        dados_completos = []
        for cat, df_cat in st.session_state["df_custos_categorias"].items():
            for _, row in df_cat.iterrows():
                valor = float(row.get("MENSAL (R$)", 0.0))
                if valor > 0: dados_completos.append({"Categoria": cat, "Subitem": row["ÍTEM"], "Valor (R$)": valor})

        if not dados_completos:
            st.warning("Adicione valores na aba de Lançamentos para ver o Dashboard ganhar vida.")
        else:
            df_dash = pd.DataFrame(dados_completos)
            
            st.markdown("#### 🔎 Filtros Dinâmicos")
            c_filtro1, c_filtro2 = st.columns([2, 1])
            todas_categorias = sorted(df_dash["Categoria"].unique().tolist())
            
            categorias_selecionadas = c_filtro1.multiselect("Selecione as categorias:", options=todas_categorias, default=todas_categorias, label_visibility="collapsed")
            
            if categorias_selecionadas:
                df_filtrado = df_dash[df_dash["Categoria"].isin(categorias_selecionadas)].copy()
                total_filtrado = df_filtrado["Valor (R$)"].sum()
                c_filtro2.metric("Total (Categorias Filtradas)", f"R$ {total_filtrado:,.2f}")
                st.write("")

                col_graf1, col_graf2 = st.columns([1, 1.2])
                
                with col_graf1:
                    st.markdown("<h5 style='text-align: center;'>Distribuição Macro</h5>", unsafe_allow_html=True)
                    df_agrupado_cat = df_filtrado.groupby("Categoria", as_index=False)["Valor (R$)"].sum()
                    fig_donut = px.pie(df_agrupado_cat, values='Valor (R$)', names='Categoria', hole=0.65, color_discrete_sequence=PALETA_GRAFICOS)
                    fig_donut.update_traces(textinfo='percent', hoverinfo='label+percent+value', textfont_size=14, marker=dict(line=dict(color='#FFFFFF', width=2)))
                    fig_donut.update_layout(showlegend=True, legend=dict(orientation="h", yanchor="bottom", y=-0.3, xanchor="center", x=0.5), margin=dict(t=20, b=20, l=20, r=20), annotations=[dict(text=f"<b>R$ {total_filtrado:,.0f}</b>", x=0.5, y=0.5, font_size=18, showarrow=False)])
                    st.plotly_chart(fig_donut, use_container_width=True)

                with col_graf2:
                    st.markdown("<h5 style='text-align: center;'>Detalhamento por Despesa</h5>", unsafe_allow_html=True)
                    df_bar = df_filtrado.sort_values(by="Valor (R$)", ascending=True)
                    max_val = df_bar["Valor (R$)"].max()
                    fig_bar = px.bar(df_bar, x="Valor (R$)", y="Subitem", color="Categoria", orientation="h", text="Valor (R$)", color_discrete_sequence=PALETA_GRAFICOS)
                    fig_bar.update_xaxes(range=[0, max_val * 1.35]) 
                    fig_bar.update_traces(texttemplate='<b>R$ %{text:,.2f}</b>', textposition='outside', textfont=dict(size=12))
                    fig_bar.update_layout(xaxis_title="", yaxis_title="", xaxis=dict(showgrid=True, gridcolor='#e8e8e8', zeroline=False), yaxis=dict(showgrid=False), plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)", height=max(350, len(df_bar) * 35), margin=dict(l=10, r=20, t=20, b=10), showlegend=False)
                    st.plotly_chart(fig_bar, use_container_width=True)

                with st.expander("📄 Ver Tabela Detalhada", expanded=False):
                    df_tabela = df_filtrado.sort_values(by="Valor (R$)", ascending=False).copy()
                    df_tabela["% do Total"] = (df_tabela["Valor (R$)"] / total_filtrado)
                    st.dataframe(df_tabela.style.format({"Valor (R$)": "R$ {:,.2f}", "% do Total": "{:.1%}"}), use_container_width=True, hide_index=True)

    with tab_lanc:
        st.subheader("⏱️ Configuração da Hora Clínica")
        col_p1, col_p2, col_p3 = st.columns(3)
        st.session_state["horas_diarias"] = col_p1.number_input("Horas Diárias", value=float(horas_diarias), step=0.5, format="%.1f")
        st.session_state["dias_semana"] = col_p2.number_input("Dias na semana", value=float(dias_semana), step=0.5, format="%.1f")
        st.session_state["qtd_salas"] = col_p3.number_input("Qtd de Salas", value=float(qtd_salas), step=1.0, format="%.1f")

        st.divider()
        st.subheader("📝 Preenchimento de Custos Mensais")
        if not st.session_state["df_custos_categorias"]:
            st.info("Nenhuma categoria. Adicione na aba de Categorias.")
        else:
            for i, categoria in enumerate(st.session_state["df_custos_categorias"].keys()):
                renderizar_categoria_dinamica(categoria, f"cat_dyn_{i}")

    with tab_cat:
        st.info("Crie novas categorias para organizar os lançamentos de custos fixos.")
        opcoes_cat = list(st.session_state["df_custos_categorias"].keys())
        
        tab_add_c, tab_ren_c, tab_del_c = st.tabs(["➕ Adicionar", "✏️ Renomear", "🗑️ Excluir"])
        with tab_add_c:
            c1, c2 = st.columns([2, 1])
            nova_cat = c1.text_input("Nome da nova categoria:", placeholder="Ex: 9. Despesas Extras")
            st.write("")
            if c2.button("Criar Categoria", use_container_width=True):
                if nova_cat and nova_cat not in st.session_state["df_custos_categorias"]:
                    st.session_state["df_custos_categorias"][nova_cat] = pd.DataFrame(columns=["ÍTEM", "MENSAL (R$)"])
                    st.rerun()
                    
        with tab_ren_c:
            c1, c2, c3 = st.columns([2, 2, 1])
            cat_renomear = c1.selectbox("Categoria atual:", options=opcoes_cat, key="sel_ren_cat")
            novo_nome_cat = c2.text_input("Mudar para:", value=cat_renomear if opcoes_cat else "", key="in_ren_cat")
            st.write("")
            if c3.button("Salvar Nome", use_container_width=True):
                if novo_nome_cat and novo_nome_cat != cat_renomear:
                    novo_dict = {}
                    for k, v in st.session_state["df_custos_categorias"].items():
                        if k == cat_renomear: novo_dict[novo_nome_cat] = v
                        else: novo_dict[k] = v
                    st.session_state["df_custos_categorias"] = novo_dict
                    st.rerun()

        with tab_del_c:
            c1, c2, c3 = st.columns([2, 1, 1])
            cat_remover = c1.selectbox("Selecione para remover:", options=opcoes_cat, key="sel_rem_cat")
            st.write("")
            if c2.button("🗑️ Remover Selecionada", use_container_width=True):
                if cat_remover in st.session_state["df_custos_categorias"]:
                    del st.session_state["df_custos_categorias"][cat_remover]
                    st.rerun()
            st.write("")
            if c3.button("⚠️ Excluir TODAS", type="primary", use_container_width=True):
                st.session_state["df_custos_categorias"] = {}
                st.rerun()

def render_equipamentos():
    cabecalho_padrao("REGISTRO DE EQUIPAMENTOS")
    st.warning("📌 Pendência Documentada: O método contábil exato de depreciação (vida útil vs taxa anual vs sessão) será refinado em atualizações futuras.")
    
    if st.button("💾 Salvar Alterações de Equip. na Nuvem", type="primary"):
        salvar_estado_nuvem()
        st.success("Equipamentos sincronizados com a Nuvem com sucesso!")

    st.subheader("⚙️ Gerenciar Equipamentos")
    
    tab_add, tab_edit, tab_del = st.tabs(["➕ Adicionar", "✏️ Editar", "🗑️ Excluir"])
    df_eq = st.session_state["df_lista_equipamentos"]
    opcoes_eq = df_eq["Nome do equipamento"].tolist()

    with tab_add:
        with st.form("form_add_eq", clear_on_submit=True):
            c1, c2, c3 = st.columns(3)
            n_nome = c1.text_input("Nome do Equipamento")
            n_valor = c2.number_input("Valor Aquisição (R$)", min_value=0.0, step=100.0, format="%.2f")
            n_vida = c3.number_input("Vida Útil (Anos)", min_value=1.0, step=1.0, format="%.1f")
            n_cap = c1.number_input("Capacidade de Aplicações / dia (R$)", min_value=0.0, step=10.0, format="%.2f")
            n_apps = c2.number_input("Aplicações (média diária)", min_value=1.0, step=1.0, format="%.1f")
            n_manut = c3.number_input("Manutenção Anual (R$)", min_value=0.0, step=10.0, format="%.2f")
            
            if st.form_submit_button("Adicionar à Lista"):
                if n_nome:
                    novo_df = pd.DataFrame([{
                        "Nome do equipamento": n_nome, 
                        "Valor de aquisição (R$)": float(n_valor),
                        "Tempo de vida útil (anos)": float(n_vida), 
                        "Capacidade de Aplicações / dia (R$)": float(n_cap),
                        "Aplicações (média diária)": float(n_apps), 
                        "Custo anual de manutenção (R$)": float(n_manut)
                    }])
                    st.session_state["df_lista_equipamentos"] = pd.concat([df_eq, novo_df], ignore_index=True)
                    salvar_estado_nuvem()
                    st.rerun()

    with tab_edit:
        eq_editar = st.selectbox("Selecione para editar:", options=opcoes_eq, key="sel_edit_eq")
        if eq_editar and not df_eq.empty:
            row_eq = df_eq[df_eq["Nome do equipamento"] == eq_editar].iloc[0]
            
            c1, c2, c3 = st.columns(3)
            novo_nome_eq = c1.text_input("Nome do Equipamento", value=row_eq["Nome do equipamento"], key="in_ren_eq")
            novo_valor_eq = c2.number_input("Valor Aquisição (R$)", value=float(row_eq["Valor de aquisição (R$)"]), min_value=0.0, step=100.0, format="%.2f", key="edit_val_eq")
            novo_vida_eq = c3.number_input("Vida Útil (Anos)", value=float(row_eq["Tempo de vida útil (anos)"]), min_value=1.0, step=1.0, format="%.1f", key="edit_vida_eq")
            novo_cap_eq = c1.number_input("Capacidade Aplicações / dia (R$)", value=float(row_eq["Capacidade de Aplicações / dia (R$)"]), min_value=0.0, step=10.0, format="%.2f", key="edit_cap_eq")
            novo_apps_eq = c2.number_input("Aplicações (média diária)", value=float(row_eq["Aplicações (média diária)"]), min_value=1.0, step=1.0, format="%.1f", key="edit_apps_eq")
            novo_manut_eq = c3.number_input("Manutenção Anual (R$)", value=float(row_eq["Custo anual de manutenção (R$)"]), min_value=0.0, step=10.0, format="%.2f", key="edit_manut_eq")
            
            st.write("")
            if st.button("💾 Salvar Edição", key="btn_salvar_edit_eq", use_container_width=True):
                if novo_nome_eq:
                    idx = df_eq.index[df_eq["Nome do equipamento"] == eq_editar].tolist()[0]
                    df_eq.at[idx, "Nome do equipamento"] = novo_nome_eq
                    df_eq.at[idx, "Valor de aquisição (R$)"] = novo_valor_eq
                    df_eq.at[idx, "Tempo de vida útil (anos)"] = novo_vida_eq
                    df_eq.at[idx, "Capacidade de Aplicações / dia (R$)"] = novo_cap_eq
                    df_eq.at[idx, "Aplicações (média diária)"] = novo_apps_eq
                    df_eq.at[idx, "Custo anual de manutenção (R$)"] = novo_manut_eq
                    
                    st.session_state["df_lista_equipamentos"] = df_eq
                    salvar_estado_nuvem()
                    st.rerun()

    with tab_del:
        c1, c2, c3 = st.columns([2, 1, 1])
        eq_remover = c1.selectbox("Selecione para remover:", options=opcoes_eq, key="sel_rem_eq")
        st.write("")
        if c2.button("🗑️ Remover Selecionado", use_container_width=True):
            st.session_state["df_lista_equipamentos"] = df_eq[df_eq["Nome do equipamento"] != eq_remover]
            salvar_estado_nuvem()
            st.rerun()
        st.write("")
        if c3.button("⚠️ Excluir TODOS", type="primary", use_container_width=True):
            st.session_state["df_lista_equipamentos"] = pd.DataFrame(columns=df_eq.columns)
            salvar_estado_nuvem()
            st.rerun()

    st.divider()
    dias_uteis = st.number_input("Dias úteis no mês para cálculo:", min_value=1.0, value=st.session_state["dias_uteis_eq"], step=1.0, format="%.0f")
    st.session_state["dias_uteis_eq"] = dias_uteis

    df_calc = st.session_state["df_lista_equipamentos"].copy()
    if "Capacidade aplicações/dia" in df_calc.columns:
        df_calc.rename(columns={"Capacidade aplicações/dia": "Capacidade de Aplicações / dia (R$)"}, inplace=True)
        st.session_state["df_lista_equipamentos"] = df_calc.copy()

    if not df_calc.empty:
        df_calc["Montante Investido"] = df_calc["Valor de aquisição (R$)"] + (df_calc["Tempo de vida útil (anos)"] * df_calc.get("Custo anual de manutenção (R$)", 0.0))
        df_calc["Depreciação Mensal"] = df_calc.apply(lambda row: row["Montante Investido"] / (row.get("Tempo de vida útil (anos)", 1) * 12) if row.get("Tempo de vida útil (anos)", 0) > 0 else 0, axis=1)
        df_calc["Custo Seção"] = df_calc.apply(lambda row: row["Depreciação Mensal"] / (row.get("Aplicações (média diária)", 1) * dias_uteis) if row.get("Aplicações (média diária)", 0) > 0 else 0, axis=1)
        
        formato_tabela = {
            "Valor de aquisição (R$)": "R$ {:,.2f}", "Capacidade de Aplicações / dia (R$)": "R$ {:,.2f}",
            "Custo anual de manutenção (R$)": "R$ {:,.2f}", "Montante Investido": "R$ {:,.2f}",
            "Depreciação Mensal": "R$ {:,.2f}", "Custo Seção": "R$ {:,.2f}"
        }
        st.dataframe(df_calc.style.format(formato_tabela, precision=2), use_container_width=True, hide_index=True)

def render_insumos():
    cabecalho_padrao("LISTA DE INSUMOS E MATERIAIS")
    
    if st.button("💾 Salvar Alterações de Insumos na Nuvem", type="primary"):
        salvar_estado_nuvem()
        st.success("Insumos sincronizados com a Nuvem com sucesso!")

    st.subheader("⚙️ Gerenciar Insumos")
    
    df_ins = st.session_state.get("df_lista_insumos", pd.DataFrame())
    if "Material" not in df_ins.columns:
        df_ins = pd.DataFrame(columns=["Material", "qt", "valor"])
        st.session_state["df_lista_insumos"] = df_ins

    tab_add, tab_edit, tab_del = st.tabs(["➕ Adicionar", "✏️ Editar", "🗑️ Excluir"])
    opcoes_ins = df_ins["Material"].tolist()

    with tab_add:
        with st.form("form_add_ins_mod", clear_on_submit=True):
            c1, c2, c3, c4 = st.columns([4, 2, 2, 2])
            n_mat = c1.text_input("Nome do Material")
            n_qt = c2.number_input("Quantidade (Embalagem)", min_value=1, step=1)
            n_val = c3.number_input("Valor (R$)", min_value=0.0, step=0.50, format="%.2f")
            
            if c4.form_submit_button("Adicionar"):
                if n_mat:
                    novo_df = pd.DataFrame([{"Material": n_mat, "qt": int(n_qt), "valor": float(n_val)}])
                    st.session_state["df_lista_insumos"] = pd.concat([df_ins, novo_df], ignore_index=True)
                    salvar_estado_nuvem()
                    st.rerun()

    with tab_edit:
        ins_editar = st.selectbox("Selecione para editar:", options=opcoes_ins, key="sel_edit_ins")
        if ins_editar and not df_ins.empty:
            row_ins = df_ins[df_ins["Material"] == ins_editar].iloc[0]
            
            c1, c2, c3 = st.columns([3, 2, 2])
            novo_nome_ins = c1.text_input("Nome do Material", value=row_ins["Material"], key="in_ren_ins")
            nova_qt_ins = c2.number_input("Quantidade (Embalagem)", value=int(row_ins["qt"]), min_value=1, step=1, key="edit_qt_ins")
            novo_val_ins = c3.number_input("Valor (R$)", value=float(row_ins["valor"]), min_value=0.0, step=0.50, format="%.2f", key="edit_val_ins")
            
            st.write("")
            if st.button("💾 Salvar Edição", key="btn_salvar_edit_ins", use_container_width=True):
                if novo_nome_ins:
                    idx = df_ins.index[df_ins["Material"] == ins_editar].tolist()[0]
                    df_ins.at[idx, "Material"] = novo_nome_ins
                    df_ins.at[idx, "qt"] = int(nova_qt_ins)
                    df_ins.at[idx, "valor"] = float(novo_val_ins)
                    
                    st.session_state["df_lista_insumos"] = df_ins
                    salvar_estado_nuvem()
                    st.rerun()

    with tab_del:
        c1, c2, c3 = st.columns([2, 1, 1])
        ins_remover = c1.selectbox("Selecione para remover:", options=opcoes_ins, key="sel_rem_ins")
        st.write("")
        if c2.button("🗑️ Remover Selecionado", use_container_width=True):
            st.session_state["df_lista_insumos"] = df_ins[df_ins["Material"] != ins_remover]
            salvar_estado_nuvem()
            st.rerun()
        st.write("")
        if c3.button("⚠️ Excluir TODOS", type="primary", use_container_width=True):
            st.session_state["df_lista_insumos"] = pd.DataFrame(columns=["Material", "qt", "valor"])
            salvar_estado_nuvem()
            st.rerun()

    st.divider()
    if not df_ins.empty:
        df_show = df_ins.copy()
        df_show["Valor Unitário (R$)"] = df_show["valor"] / df_show["qt"]
        st.dataframe(df_show.style.format({"qt": "{:.0f}", "valor": "R$ {:.2f}", "Valor Unitário (R$)": "R$ {:.2f}"}), use_container_width=True, hide_index=True)

def render_taxas():
    cabecalho_padrao("IMPOSTOS E TAXAS")
    
    if st.button("💾 Salvar Alterações de Taxas na Nuvem", type="primary"):
        salvar_estado_nuvem()
        st.success("Taxas sincronizadas com a Nuvem com sucesso!")

    st.subheader("⚙️ Gerenciar Taxas")
    
    tab_add, tab_edit, tab_del = st.tabs(["➕ Adicionar", "✏️ Editar", "🗑️ Excluir"])
    df_taxas = st.session_state["df_lista_taxas"]
    opcoes_taxas = df_taxas["Taxa"].tolist()

    with tab_add:
        with st.form("form_add_taxa_mod", clear_on_submit=True):
            c1, c2, c3 = st.columns([4, 2, 2])
            n_taxa = c1.text_input("Nome da Taxa (Ex: Crédito 2x)")
            n_pct = c2.number_input("Porcentagem (%)", min_value=0.0, step=0.10, format="%.2f")
            if c3.form_submit_button("Adicionar"):
                if n_taxa:
                    novo_df = pd.DataFrame([{"Taxa": n_taxa, "Porcentagem (%)": float(n_pct)}])
                    st.session_state["df_lista_taxas"] = pd.concat([df_taxas, novo_df], ignore_index=True)
                    salvar_estado_nuvem()
                    st.rerun()

    with tab_edit:
        taxa_editar = st.selectbox("Selecione para editar:", options=opcoes_taxas, key="sel_edit_taxa")
        if taxa_editar and not df_taxas.empty:
            row_taxa = df_taxas[df_taxas["Taxa"] == taxa_editar].iloc[0]
            
            c1, c2 = st.columns([3, 2])
            novo_nome_taxa = c1.text_input("Nome da Taxa", value=row_taxa["Taxa"], key="in_ren_taxa")
            novo_pct_taxa = c2.number_input("Porcentagem (%)", value=float(row_taxa["Porcentagem (%)"]), min_value=0.0, step=0.10, format="%.2f", key="edit_pct_taxa")
            
            st.write("")
            if st.button("💾 Salvar Edição", key="btn_salvar_edit_taxa", use_container_width=True):
                if novo_nome_taxa:
                    idx = df_taxas.index[df_taxas["Taxa"] == taxa_editar].tolist()[0]
                    df_taxas.at[idx, "Taxa"] = novo_nome_taxa
                    df_taxas.at[idx, "Porcentagem (%)"] = novo_pct_taxa
                    
                    st.session_state["df_lista_taxas"] = df_taxas
                    salvar_estado_nuvem()
                    st.rerun()

    with tab_del:
        c1, c2, c3 = st.columns([2, 1, 1])
        taxa_remover = c1.selectbox("Remover Taxa:", options=opcoes_taxas, key="sel_rem_taxa")
        st.write("")
        if c2.button("🗑️ Remover", use_container_width=True):
            st.session_state["df_lista_taxas"] = df_taxas[df_taxas["Taxa"] != taxa_remover]
            salvar_estado_nuvem()
            st.rerun()
        st.write("")
        if c3.button("⚠️ Excluir TODAS", type="primary", use_container_width=True):
            st.session_state["df_lista_taxas"] = pd.DataFrame(columns=df_taxas.columns)
            salvar_estado_nuvem()
            st.rerun()

    st.divider()
    if not df_taxas.empty:
        st.dataframe(df_taxas.style.format({"Porcentagem (%)": "{:.2f}%"}), use_container_width=True, hide_index=True)

def render_protocolos():
    cabecalho_padrao("🚀 PROTOCOLOS E JORNADAS")
    
    if "protocolo_atual" not in st.session_state:
        st.session_state["protocolo_atual"] = {"nome": "", "descricao": "", "beneficio": "", "itens": []}
    
    # --- 1. CONFIGURAÇÃO COMERCIAL ---
    st.subheader("1. Identidade do Protocolo")
    c_p1, c_p2 = st.columns(2)
    st.session_state.protocolo_atual["nome"] = c_p1.text_input("Nome do Protocolo (Ex: Jornada Bariátrica)", value=st.session_state.protocolo_atual["nome"])
    st.session_state.protocolo_atual["descricao"] = st.text_area("Descrição da Jornada", value=st.session_state.protocolo_atual.get("descricao", ""), placeholder="O que está incluído no acompanhamento...")
    st.session_state.protocolo_atual["beneficio"] = st.text_input("Principal Benefício / Proposta de Valor", value=st.session_state.protocolo_atual.get("beneficio", ""))

    # --- 2. MONTAGEM DO PACOTE ---
    st.divider()
    st.subheader("2. Composição da Jornada")
    
    db_servicos = st.session_state.get("db_servicos", {})
    if not db_servicos:
        st.warning("Cadastre serviços na Ficha Técnica primeiro.")
        return

    with st.expander("➕ Adicionar Serviço ao Pacote", expanded=True):
        with st.form("form_add_item_prot", clear_on_submit=True):
            c_a1, c_a2, c_a3 = st.columns([3,1,1])
            serv_nome = c_a1.selectbox("Serviço", list(db_servicos.keys()))
            qtd = c_a2.number_input("Qtde (Sessões)", min_value=1, value=1)
            desc_lin = c_a3.number_input("Desc. Linha (%)", min_value=0.0, max_value=100.0, value=0.0)
            
            if st.form_submit_button("ADICIONAR ITEM"):
                st.session_state.protocolo_atual["itens"].append({
                    "servico": serv_nome,
                    "qtd": qtd,
                    "desconto": desc_lin
                })
                st.rerun()

    # --- 3. CÁLCULO INDIVIDUALIZADO (Lógica da Planilha) ---
    itens = st.session_state.protocolo_atual["itens"]
    if itens:
        dados_calculados = []
        valor_hora_global = st.session_state.get("valor_hora", 48.14)
        usa_ind = st.session_state.get("indireto") == "Sim"

        df_taxas_g = st.session_state.get("df_lista_taxas", pd.DataFrame())

        for i, item in enumerate(itens):
            ficha = db_servicos[item["servico"]]
            
            # Preço Base
            p_un = ficha.get("preco_escolhido", 0.0)
            p_tot_bruto = p_un * item["qtd"]
            p_tot_com_desc = p_tot_bruto * (1 - (item["desconto"]/100))
            
            # Taxas Individuais
            txs = ficha.get("taxas", {})
            t_imp = txs.get("aliquota_imposto", 6.0)
            t_com = txs.get("comissao", 0.0)
            
            cenario_s_l = txs.get("cenario_cartao", "Crédito 1x")
            t_car = 0.0
            if not df_taxas_g.empty and cenario_s_l in df_taxas_g["Taxa"].values:
                t_car = float(df_taxas_g[df_taxas_g["Taxa"] == cenario_s_l]["Porcentagem (%)"].iloc[0])
            
            v_deducoes = p_tot_com_desc * ((t_imp + t_com + t_car)/100)
            resultado_liquido = p_tot_com_desc - v_deducoes
            
            # Repasse Médico individual da ficha
            p_rep_med = ficha.get("repasse_percentual", 0.0)
            v_rep_med = resultado_liquido * (p_rep_med/100)
            
            # Custos Operacionais
            t_min = ficha.get("tempo_min", 60.0)
            c_exec = (t_min/60) * valor_hora_global if usa_ind else 0.0
            
            df_maq_f = pd.DataFrame(ficha.get("maquinas", []))
            c_maq = pd.to_numeric(df_maq_f["custo"], errors="coerce").fillna(0.0).sum() if not df_maq_f.empty else 0.0
            
            c_alu_h = ficha.get("custo_aluguel", 0.0)
            c_alu = c_alu_h * (t_min/60)
            
            df_ins_f = pd.DataFrame(ficha.get("insumos", []))
            c_ins = 0.0
            if not df_ins_f.empty:
                c_ins = (pd.to_numeric(df_ins_f["QT"], errors="coerce").fillna(0.0) * pd.to_numeric(df_ins_f["Preço (R$)"], errors="coerce").fillna(0.0)).sum()
            
            df_o_f = pd.DataFrame(ficha.get("outros_custos", []))
            c_outros_s = pd.to_numeric(df_o_f["Valor (R$)"], errors="coerce").fillna(0.0).sum() if not df_o_f.empty else 0.0

            c_total_operacional = (c_exec + c_maq + c_alu + c_ins + c_outros_s + ficha.get("repasse_fixo", 0.0)) * item["qtd"]
            
            lucro_item = resultado_liquido - v_rep_med - c_total_operacional
            
            dados_calculados.append({
                "Serviço": item["servico"],
                "Qtde": item["qtd"],
                "Venda (R$)": p_tot_com_desc,
                "Resultado Líquido": resultado_liquido,
                "Repasse Médico": v_rep_med,
                "Custo Serviço": c_total_operacional,
                "Lucro": lucro_item,
                "% Lucro": (lucro_item / p_tot_com_desc) if p_tot_com_desc > 0 else 0.0
            })

        df_prot = pd.DataFrame(dados_calculados)
        
        # --- EXIBIÇÃO TABELA ROXA (Igual Planilha) ---
        st.markdown("**Tabela de Resultados por Item:**")
        st.dataframe(df_prot.style.format({
            "Venda (R$)": "R$ {:,.2f}", "Resultado Líquido": "R$ {:,.2f}", 
            "Repasse Médico": "R$ {:,.2f}", "Custo Serviço": "R$ {:,.2f}", 
            "Lucro": "R$ {:,.2f}", "% Lucro": "{:.1%}"
        }), use_container_width=True, hide_index=True)

        if st.button("🗑️ Limpar Pacote"):
            st.session_state.protocolo_atual["itens"] = []
            st.rerun()

        # --- 4. RESUMO DE REPASSES ---
        st.divider()
        col_res1, col_res2 = st.columns(2)
        
        with col_res1:
            st.markdown("### 💰 Consolidação do Protocolo")
            venda_tot = df_prot["Venda (R$)"].sum()
            custo_tot = df_prot["Custo Serviço"].sum()
            repasse_tot = df_prot["Repasse Médico"].sum()
            lucro_tot = df_prot["Lucro"].sum()
            
            st.metric("PREÇO FINAL DO PROTOCOLO", f"R$ {venda_tot:,.2f}")
            st.metric("LUCRO LÍQUIDO FINAL (CLÍNICA)", f"R$ {lucro_tot:,.2f}", delta=f"{lucro_tot/venda_tot:.1%}" if venda_tot > 0 else "0%")
            
        with col_res2:
            st.markdown("### 👥 Valor a Distribuir (Repasses)")
            df_repasses = df_prot.groupby("Serviço")["Repasse Médico"].sum().reset_index(name="Total Repasse")
            st.dataframe(df_repasses.style.format({"Total Repasse": "R$ {:,.2f}"}), use_container_width=True, hide_index=True)
            st.metric("Total Pago à Equipe", f"R$ {repasse_tot:,.2f}")

        # --- 5. GERAÇÃO DE PROPOSTA ---
        def gerar_word_proposta(nome_pacote, desc, benef, df_itens, total_venda):
            doc = Document()
            doc.add_heading(f"Proposta Comercial: {nome_pacote}", 0)
            if desc: doc.add_paragraph(f"Descrição: {desc}\n")
            if benef: doc.add_paragraph(f"Benefícios: {benef}\n")
            
            doc.add_paragraph("Detalhes dos serviços inclusos no seu pacote personalizado.\n")
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Serviço'
            hdr_cells[1].text = 'Quantidade'
            hdr_cells[2].text = 'Subtotal (R$)'
            
            for index, row in df_itens.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['Serviço'])
                row_cells[1].text = f"{row['Qtde']} sessões"
                row_cells[2].text = f"R$ {row['Venda (R$)']:,.2f}"
                
            p = doc.add_paragraph("\n")
            p.add_run(f"Valor Total do Investimento: R$ {total_venda:,.2f}").bold = True
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        st.write("")
        c_save1, c_save2 = st.columns(2)
        with c_save1:
            if st.button("💾 SALVAR PROTOCOLO NA NUVEM", type="primary", use_container_width=True):
                if "protocolos_db" not in st.session_state: st.session_state.protocolos_db = []
                st.session_state.protocolos_db.append(st.session_state.protocolo_atual)
                salvar_estado_nuvem()
                st.success("Protocolo arquivado com sucesso!")
        with c_save2:
            arquivo_word = gerar_word_proposta(st.session_state.protocolo_atual["nome"], st.session_state.protocolo_atual["descricao"], st.session_state.protocolo_atual["beneficio"], df_prot, venda_tot)
            st.download_button(
                label="📄 Gerar Proposta (Word)",
                data=arquivo_word,
                file_name=f"Proposta_{st.session_state.protocolo_atual['nome']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

    else:
        st.info("Adicione itens acima para começar a montar o protocolo.")

# ==========================================
# ROTEAMENTO
# ==========================================
if modulo_selecionado == "0. Início (Onboarding)": render_onboarding()
elif modulo_selecionado == "1. Ficha Técnica (Precificação)": render_ficha_tecnica()
elif modulo_selecionado == "2. Custos Fixos e Hora Clínica": render_custos_fixos()
elif modulo_selecionado == "3. Registro de Equipamentos": render_equipamentos()
elif modulo_selecionado == "4. Insumos e Materiais": render_insumos()
elif modulo_selecionado == "5. Impostos e Taxas": render_taxas()
elif modulo_selecionado == "6. Protocolos (Jornadas) 🚀": render_protocolos()